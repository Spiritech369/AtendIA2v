from __future__ import annotations

from io import BytesIO
from uuid import uuid4

import fitz  # type: ignore[import-untyped]
import pytest
from docx import Document
from openpyxl import Workbook

from atendia.knowledge.os import InMemoryKnowledgeRepository
from atendia.knowledge.os.ingestion import KnowledgeIngestionService
from atendia.knowledge.os.retrieval import KnowledgeRetrievalService


def _pdf_bytes(text: str) -> bytes:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), text)
    data = document.tobytes()
    document.close()
    return data


def _docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Support Policy", level=1)
    document.add_paragraph("Premium support is available Monday morning.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "service"
    table.cell(0, 1).text = "price"
    table.cell(1, 0).text = "Priority onboarding"
    table.cell(1, 1).text = "$299"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _xlsx_bytes() -> bytes:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Pricing"
    worksheet.append(["service", "category", "price", "currency"])
    worksheet.append(["Premium Wash", "detailing", 199, "USD"])
    worksheet.append(["Basic Wash", "detailing", 49, "USD"])
    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


@pytest.mark.asyncio
async def test_pdf_ingestion_retrieval_includes_page_citation():
    tenant_id = uuid4()
    repo = InMemoryKnowledgeRepository()
    ingestion = KnowledgeIngestionService(repo)
    retrieval = KnowledgeRetrievalService(repo)

    source, items, chunks = await ingestion.ingest_file(
        tenant_id=tenant_id,
        name="Refund policy",
        filename="refund-policy.pdf",
        data=_pdf_bytes("Refunds are available within 14 days on page one."),
        content_type="policy",
    )

    evidence = await retrieval.retrieve(tenant_id=tenant_id, query="refunds 14 days")

    assert source.status == "active"
    assert items
    assert chunks
    assert evidence.answerable is True
    assert evidence.citations[0].metadata["page"] == 1
    assert evidence.citations[0].metadata["file_type"] == "pdf"


@pytest.mark.asyncio
async def test_docx_ingestion_retrieves_text_and_table_rows():
    tenant_id = uuid4()
    repo = InMemoryKnowledgeRepository()
    ingestion = KnowledgeIngestionService(repo)
    retrieval = KnowledgeRetrievalService(repo)

    source, items, _chunks = await ingestion.ingest_file(
        tenant_id=tenant_id,
        name="Support guide",
        filename="support-guide.docx",
        data=_docx_bytes(),
        content_type="services",
    )

    evidence = await retrieval.retrieve(tenant_id=tenant_id, query="Priority onboarding 299")

    assert source.status == "active"
    assert any(item.structured_data for item in items)
    assert evidence.answerable is True
    assert evidence.citations[0].metadata["file_type"] == "docx"
    assert evidence.citations[0].metadata["table"] == 1


@pytest.mark.asyncio
async def test_csv_catalog_preserves_rows_and_retrieves_price_terms():
    tenant_id = uuid4()
    repo = InMemoryKnowledgeRepository()
    ingestion = KnowledgeIngestionService(repo)
    retrieval = KnowledgeRetrievalService(repo)

    csv_data = (
        b"service,category,price,currency\n"
        b"Premium Wash,detailing,199,USD\n"
        b"Basic Wash,detailing,49,USD\n"
    )
    source, items, _chunks = await ingestion.ingest_file(
        tenant_id=tenant_id,
        name="Service catalog",
        filename="catalog.csv",
        data=csv_data,
        content_type="pricing",
    )

    evidence = await retrieval.retrieve(tenant_id=tenant_id, query="Premium Wash 199")

    assert source.type == "table"
    assert any(item.structured_data and item.structured_data["price"] == "199" for item in items)
    assert evidence.answerable is True
    assert evidence.citations[0].metadata["row_index"] == 1
    assert "price: 199" in evidence.citations[0].snippet


@pytest.mark.asyncio
async def test_xlsx_catalog_retrieval_includes_sheet_citation():
    tenant_id = uuid4()
    repo = InMemoryKnowledgeRepository()
    ingestion = KnowledgeIngestionService(repo)
    retrieval = KnowledgeRetrievalService(repo)

    source, items, _chunks = await ingestion.ingest_file(
        tenant_id=tenant_id,
        name="Workbook catalog",
        filename="catalog.xlsx",
        data=_xlsx_bytes(),
        content_type="pricing",
    )

    evidence = await retrieval.retrieve(tenant_id=tenant_id, query="Basic Wash 49")

    assert source.type == "table"
    assert any(item.metadata.get("sheet") == "Pricing" for item in items)
    assert evidence.answerable is True
    assert evidence.citations[0].metadata["sheet"] == "Pricing"


@pytest.mark.asyncio
async def test_file_retrieval_does_not_mix_tenants():
    tenant_a = uuid4()
    tenant_b = uuid4()
    repo = InMemoryKnowledgeRepository()
    ingestion = KnowledgeIngestionService(repo)
    retrieval = KnowledgeRetrievalService(repo)

    await ingestion.ingest_file(
        tenant_id=tenant_a,
        name="Tenant A catalog",
        filename="catalog.csv",
        data=b"service,price\nPrivate consultation,150\n",
        content_type="pricing",
    )
    await ingestion.ingest_file(
        tenant_id=tenant_b,
        name="Tenant B catalog",
        filename="catalog.csv",
        data=b"service,price\nPrivate consultation,900\n",
        content_type="pricing",
    )

    evidence = await retrieval.retrieve(tenant_id=tenant_a, query="Private consultation")

    assert evidence.answerable is True
    assert all(citation.source_name == "Tenant A catalog" for citation in evidence.citations)
    assert "900" not in " ".join(citation.snippet for citation in evidence.citations)


@pytest.mark.asyncio
async def test_image_ingestion_is_metadata_only_stub():
    tenant_id = uuid4()
    repo = InMemoryKnowledgeRepository()
    ingestion = KnowledgeIngestionService(repo)

    source, items, chunks = await ingestion.ingest_file(
        tenant_id=tenant_id,
        name="Reference image",
        filename="menu.png",
        data=b"not-a-real-image",
        content_type="catalog",
    )

    assert source.status == "partially_processed"
    assert source.metadata["image_ingestion"] == "metadata_only"
    assert source.metadata["warnings"] == ["image_ocr_not_enabled"]
    assert items == []
    assert chunks == []
