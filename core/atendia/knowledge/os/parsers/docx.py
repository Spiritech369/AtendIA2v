from __future__ import annotations

from io import BytesIO

from docx import Document

from atendia.knowledge.os.parsers.base import (
    ParsedDocument,
    ParsedSection,
    ParsedTable,
    normalize_text,
    stringify_cell,
)


def parse_docx(data: bytes, *, filename: str = "") -> ParsedDocument:
    document = Document(BytesIO(data))
    sections: list[ParsedSection] = []
    tables: list[ParsedTable] = []
    current_heading = filename or "DOCX document"
    current_lines: list[str] = []
    section_index = 1

    def flush_section() -> None:
        nonlocal current_lines, section_index
        text = normalize_text("\n".join(current_lines))
        if not text:
            current_lines = []
            return
        sections.append(
            ParsedSection(
                text=text,
                title=current_heading,
                metadata={
                    "file_type": "docx",
                    "section": section_index,
                    "heading": current_heading,
                },
            )
        )
        section_index += 1
        current_lines = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = (paragraph.style.name if paragraph.style is not None else "").casefold()
        if style.startswith("heading"):
            flush_section()
            current_heading = text
        else:
            current_lines.append(text)
    flush_section()

    for table_index, table in enumerate(document.tables, start=1):
        raw_rows = [
            [stringify_cell(cell.text) for cell in row.cells]
            for row in table.rows
            if any(stringify_cell(cell.text) for cell in row.cells)
        ]
        if not raw_rows:
            continue
        headers = raw_rows[0]
        if not any(headers):
            headers = [f"column_{index + 1}" for index in range(len(raw_rows[0]))]
            data_rows = raw_rows
        else:
            headers = [header or f"column_{index + 1}" for index, header in enumerate(headers)]
            data_rows = raw_rows[1:]
        rows = [
            {
                headers[col_index]: row[col_index] if col_index < len(row) else ""
                for col_index in range(len(headers))
            }
            for row in data_rows
        ]
        tables.append(
            ParsedTable(
                headers=headers,
                rows=rows,
                reference={"table": table_index},
                metadata={"file_type": "docx"},
            )
        )

    table_lines = []
    for table in tables:
        for index, row in enumerate(table.rows, start=1):
            table_lines.append(
                f"table {table.reference.get('table')} row {index}: "
                + "; ".join(f"{key}: {value}" for key, value in row.items() if value)
            )
    extracted = normalize_text(
        "\n".join([*(section.text for section in sections), *table_lines])
    )
    return ParsedDocument(
        extracted_text=extracted,
        sections=sections,
        tables=tables,
        metadata={
            "filename": filename,
            "file_type": "docx",
            "section_count": len(sections),
            "table_count": len(tables),
        },
        warnings=[],
    )
